from django import forms
import requests
from django.http import HttpResponseRedirect
from django.shortcuts import render
# from .forms import UserForm
from docx import Document
from docx2pdf import convert
import urllib.request
import shutil


def replace_string(lastname, name, patronymic):
    changes = lastname + ' ' + name + ' ' + patronymic
    print(changes)
    dict_men = ('Шпаков Роман Сергеевич', 'Иванов')
    if changes in dict_men:
        print('jr')
        doc = Document('main\ату.docx')
        # doc = Document('ату.docx')
        for p in doc.paragraphs:
            if 'LAST' in p.text:
                print('1')
                inline = p.runs
                # Loop added to work with runs (strings with same style)
                for i in range(len(inline)):
                    if 'LAST' in inline[i].text:
                        print('2')
                        text = inline[i].text.replace('LAST', changes)
                        inline[i].text = text
                print(p.text)

        doc.save('main\dest1.docx')
        # doc.save('dest1.docx')
        convert('main\dest1.docx', 'main\dest1.pdf')
        return 1
    else:
        print('no')
# replace_string('Шпаков')
# def downloads():
#
#     url = "C:\\pythonProject\\main\\dest1.pdf"
#     output_file = "dest1.pdf"
#     with urllib.request.urlopen(url) as response, open(output_file, 'wb') as out_file:
#         shutil.copyfileobj(response, out_file)